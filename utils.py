from docx import Document
import fitz  # PyMuPDF
import os
import PyPDF2


def write_string_to_word(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

def read_docx(file):
    doc = Document(file)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    return text

def read_pdf(uploaded_file):
    text = ''
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    num_pages = len(pdf_reader.pages)

    for page_number in range(num_pages):
        page = pdf_reader.pages[page_number]
        text += page.extract_text()

    return text
